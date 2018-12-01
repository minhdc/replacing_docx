import os
import docx
import constant
'''
    1. dictionary
        - encoding
        - file
    2. replacing
'''

#load dictionary
def get_dict_from_file(dict_file_as_docx_obj):
    dictionary_as_dict = {}
    for key_value in dict_file_as_docx_obj.paragraphs:    
        #key_value.text.replace(' ','')
        key = key_value.text.split(':')[0]
        value = key_value.text.split(':')[1]
        dictionary_as_dict.update({key:value})
    return dictionary_as_dict

#need to use regex 
def docx_paragraphs_to_string(docx_paragraphs):
    result = ""
    for paragraph in docx_paragraphs:
        paragraph = paragraph.text.replace('\t',' ')
        paragraph = paragraph.replace('\n',' ')
        result += paragraph
    return result
    
def do_replace(input_string,single_dict):
    for k,v in single_dict.items():
        if k in input_string:
            input_string = input_string.replace(k,v)
    return input_string

#load file
#input_file = docx.Document('example.docx')
#input_dictionary = docx.Document('dictionary.docx')

result = ""
list_of_dicts = [{},{},{}]
list_of_files = []

#1.load all dicts into a list 
for dict_name in os.listdir(constant.DICT_LOC):
    if '3' in dict_name:
        list_of_dicts[0] = get_dict_from_file(docx.Document(constant.DICT_LOC+os.sep+dict_name))
    elif '2' in dict_name:
        list_of_dicts[1] = get_dict_from_file(docx.Document(constant.DICT_LOC+os.sep+dict_name))
    else:
        list_of_dicts[2] = get_dict_from_file(docx.Document(constant.DICT_LOC+os.sep+dict_name))       

for e in list_of_dicts:
    if e:
        print(e)

#2.load all input files
for f_name in os.listdir(constant.INPUT_LOC):
    list_of_files.append(f_name)

for each_file in list_of_files:
    original = docx_paragraphs_to_string(docx.Document(constant.INPUT_LOC+os.sep+each_file).paragraphs)
    result = docx_paragraphs_to_string(docx.Document(constant.INPUT_LOC+os.sep+each_file).paragraphs)    
    
    for each_dict in list_of_dicts:
        if each_dict:
            result = do_replace(result,each_dict)             

    result = result.split(' ')
    result = [result[i:i+constant.EXPECTED_LINE_LENGTH] for i in range(0, len(result),constant.EXPECTED_LINE_LENGTH)]
    original = original.split(' ')
    original = [original[i:i+constant.EXPECTED_LINE_LENGTH] for i in range(0,len(original),constant.EXPECTED_LINE_LENGTH)]

    print(result)
    print(original)

    replaced_example = docx.Document()
    for i in range(0,len(original)):        
        replaced_example.add_paragraph(' '.join(original[i])) 
        replaced_example.add_paragraph(' '.join(result[i])) 
    replaced_example.save(constant.OUTPUT_LOC+os.sep+"replaced "+each_file)
    #print original

    #print replaced

    